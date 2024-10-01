# import threading

import smtplib
from email.message import EmailMessage
import time
import os

def treatment_p(id, fio, contacts,type_service, service, date, category):

    if category == None:
        category = ''

    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    # Создайте новый документ
    doc = Document()

    # Функция для добавления текста по центру
    def add_centered_text(doc, text, bold=False, font_size=12):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Первая строка
    add_centered_text(doc, "ЗАЯВКА № ______", bold=True, font_size=14)

    date = str(date).split('-')

    # # Получение даты
    # date_str = date.strftime("от «%d»_%m_%Y г.")

    # Вторая строка
    add_centered_text(doc, f'от {date[2]}_{date[1]}_{date[0]} г.', font_size=12)

    # Третья строка
    add_centered_text(doc, "на оказание услуги по выездному обслуживанию заявителей", font_size=12)

    # Добавление таблицы
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'

    # Первая строка таблицы
    cell = table.cell(0, 0)
    cell.merge(table.cell(0, 1))
    cell_paragraph = cell.paragraphs[0]
    cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = cell_paragraph.add_run("Сведения о заказчике")
    run.bold = True

    # Вторая строка таблицы
    table.cell(1, 0).text = "Фамилия Имя Отчество/ Наименование заказчика"

    # Вторая строка таблицы
    table.cell(1, 1).text = fio

    # Третья строка таблицы
    table.cell(2, 0).text = "Контактный телефон/ адрес проживания заказчика адрес электронной почты"

    # Третья строка таблицы
    table.cell(2, 1).text = contacts

    # Четвёртая строка таблицы
    table.cell(3, 0).text = "Наименование категории граждан, для которых организация выезда через МФЦ осуществляется бесплатно*"

    # Четвёртая строка таблицы
    table.cell(3, 1).text = category

    # Пятая строка таблицы
    table.cell(4, 0).text = "Вид услуги"
    if type_service == '1':
        table.cell(4, 1).text = "Выезд к заявителю с целью:\n⩗ приёма документов\n□ доставки документов"
    else:
        table.cell(4, 1).text = "Выезд к заявителю с целью:\n□ приёма документов\n⩗ доставки документов"

    # Шестая строка таблицы
    table.cell(5, 0).text = "Наименование услуги"

    # Шестая строка таблицы
    table.cell(5, 1).text = service

    # После таблицы
    doc.add_paragraph()
    add_centered_text(doc, "Заявку принял: __________(подпись) / __________(ФИО специалиста)", font_size=12)

    # Последняя строка справа
    paragraph = doc.add_paragraph("* - заполняется при выездном обслуживании льготных категорий граждан")
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Сохраните документ
    modified_attachment_path = "C:\\Users\\admin\\Desktop\\mail\\zayavka_" + str(id) + ".docx"

    doc.save(modified_attachment_path)

    print("Документ успешно создан и сохранён.")

    # Заполните эти переменные значениями для своего почтового сервера и учетной записи отправителя
    smtp_server = "smtp.mfc.tomsk.ru"
    # smtp_server = "smtp.yandex.ru"
    smtp_port = 587  # Порт для TLS
    sender_email_address = "neverov@mfc.tomsk.ru"
    sender_password = "85HtnPdj"  # Убедитесь, что это пароль приложения
    # sender_email_address = "m.tosp@yandex.ru"
    # sender_password = "cwlecvijmxlpkvfo"  # Убедитесь, что это пароль приложения

    # Создайте сообщение электронной почты
    message = EmailMessage()
    message["Subject"] = f"Заявка на платную услугу от {fio}"
    message["From"] = sender_email_address
    message["To"] = 'msptosp@mfc.tomsk.ru'  # Адрес получателя

    # Установите текстовое содержимое письма
    message.set_content(f"Заявка на платную у слугу от {fio}")

    # Прикрепите файл Excel к письму
    attachment_path = "C:\\Users\\admin\\Desktop\\mail\\zayavka_" + str(id) + ".docx"
    with open(attachment_path, "rb") as attachment:
        message.add_attachment(attachment.read(), maintype="application", subtype="vnd.openxmlformats-officedocument.wordprocessingml.document", filename=os.path.basename(attachment_path))

    # Функция для отправки письма
    def send_email():
        condition = False
        try:
            # Подключитесь к SMTP-серверу
            with smtplib.SMTP(smtp_server, smtp_port) as server:
                server.starttls()  # Инициализация TLS
                server.login(sender_email_address, sender_password)
                server.send_message(message)  # Отправьте сообщение
            print("Письмо успешно отправлено.")

            # Удалите документ
            os.remove(modified_attachment_path)

            condition = True
        except smtplib.SMTPAuthenticationError as e:
            print(f"Ошибка аутентификации mail: {e}")
        except Exception as e:
            print(f"Произошла ошибка mail: {e}")

        return condition

    # Попробуйте отправить письмо несколько раз с интервалами
    attempts = 3
    for attempt in range(attempts):
        print(f"Попытка {attempt + 1} из {attempts}")

        condition = send_email()
        if condition:
            break
        elif attempt < attempts - 1:
            print("Ожидание перед повторной попыткой...")
            time.sleep(10)  # Ожидание 10 секунд перед следующей попыткой

    return

import mysql.connector

def process_mail():

    # Создание одиночного соединения
    dbconfig = {
        'host': "172.18.11.103",
        'user': "root",
        'password': "enigma1418",
        'database': "mdtomskbot",
    }

    connection = mysql.connector.connect(**dbconfig)

    # Выполнение запросов и обновление данных
    try:
        cursor = connection.cursor()

        while True:
            # global exit_event
            update_query = "SELECT * FROM application;"
            cursor.execute(update_query)
            data = cursor.fetchall()

            for row in data:
                treatment_p(row[0], row[1], row[2], row[3], row[4], row[5], row[6])
                update_query = f"DELETE FROM application WHERE id = {row[0]};"
                cursor.execute(update_query)

            connection.commit()  # Фиксация транзакции

    except Exception as e:
        # Вывод подробной информации об ошибке
        print(f"Поймано исключение: {type(e).__name__}")
        print(f"Сообщение об ошибке: {str(e)}")
        import traceback
        print("Трассировка стека (stack trace):")
        traceback.print_exc()

    finally:
        if 'cursor' in locals() or 'cursor' in globals():
            cursor.close()  # Важно закрыть курсор
        if 'connection' in locals() or 'connection' in globals():
            connection.close()  # Важно закрыть соединение после его использования

def process_file():
    # try:
    #     from email.mime.base import MIMEBase
    #     from email import encoders

    #     # Путь к папке, где находятся файлы
    #     folder_path = "C:\\Users\\admin\\Desktop\\file"

    #     # Функция для поиска первого файла, содержащего "info" в имени
    #     def find_file_with_info(folder):
    #         for filename in os.listdir(folder):
    #             if "info" in filename:
    #                 return os.path.join(folder, filename)
    #         return None

    #     # Найдите файл с "info" в имени
    #     file_with_info = find_file_with_info(folder_path)

    #     if not file_with_info:
    #         # print("Файл, содержащий 'info', не найден.")
    #         return

    #     # Функция для извлечения текста из файла
    #     def extract_text_from_file(filepath):
    #         with open(filepath, 'r', encoding='utf-8') as f:
    #                 return f.read()

    #         # import mimetypes
    #         # mime_type, _ = mimetypes.guess_type(filepath)
    #         # if mime_type == 'text/plain':  # Если это .txt файл
    #         #     with open(filepath, 'r', encoding='utf-8') as f:
    #         #         return f.read()
    #         # elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':  # Если это .docx файл
    #         #     from docx import Document
    #         #     doc = Document(filepath)
    #         #     return "\n".join([para.text for para in doc.paragraphs])
    #         # elif mime_type == 'application/pdf':  # Если это .pdf файл
    #         #     import PyPDF2
    #         #     with open(filepath, 'rb') as f:
    #         #         reader = PyPDF2.PdfReader(f)
    #         #         text = ""
    #         #         for page in reader.pages:
    #         #             text += page.extract_text()
    #         #         return text
    #         # else:
    #         #     return "Неподдерживаемый формат файла"

    #         # # Проверка расширения файла
    #         # extension = os.path.splitext(filepath)[1].lower()

    #         # if extension == '.txt':  # Если это .txt файл
    #         #     with open(filepath, 'r', encoding='utf-8') as f:
    #         #         return f.read()
    #         # elif extension == '.docx':  # Если это .docx файл
    #         #     from docx import Document
    #         #     doc = Document(filepath)
    #         #     return "\n".join([para.text for para in doc.paragraphs])
    #         # elif extension == '.pdf':  # Если это .pdf файл
    #         #     import PyPDF2
    #         #     with open(filepath, 'rb') as f:
    #         #         reader = PyPDF2.PdfReader(f)
    #         #         text = ""
    #         #         for page in reader.pages:
    #         #             text += page.extract_text()
    #         #         return text
    #         # else:
    #         #     return "Неподдерживаемый формат файла"

    #     # Извлекаем текст из файла
    #     file_text = extract_text_from_file(file_with_info)

    #     # Функция для получения первой части имени файла до символа "_"
    #     def extract_search_string_from_filename(filename):
    #         return filename.split("_")[1] if "_" in filename else None

    #     # # Функция для поиска первого файла в папке
    #     # def get_first_file_in_folder(folder):
    #     #     files = os.listdir(folder)
    #     #     if files:
    #     #         return files[0]  # Вернуть первый файл в папке
    #     #     return None

    #     # Функция для поиска первого файла, не содержащего "info" в имени
    #     def get_first_file_in_folder(folder):
    #         files = os.listdir(folder)
    #         for filename in files:
    #             if not "info" in filename:
    #                 return os.path.join(folder, filename)
    #         return None

    #     # # Функция для поиска первого файла в папке, не содержащего "info" в имени
    #     # def get_first_file_in_folder(folder):
    #     #     files = os.listdir(folder)
    #     #     for file in files:
    #     #         if "info" not in file:  # Проверка, что в имени файла нет "info"
    #     #             return file  # Вернуть первый файл, не содержащий "info"
    #     #     return None  # Вернуть None, если все файлы содержат "info" в имени

    #     # Найдите первый файл в папке
    #     first_file = get_first_file_in_folder(folder_path)

    #     if not first_file:
    #         # print("В папке нет файлов.")
    #         return

    #     # Извлекаем строку для поиска из имени файла
    #     search_string = str(extract_search_string_from_filename(first_file)).split('.')[0]

    #     if not search_string:
    #         # print(f"Не удалось извлечь строку из имени файла: {first_file}")
    #         return

    #     # # Поиск файла, который содержит строку в имени
    #     # def find_file_in_folder(folder, search_term):
    #     #     file_1 = None
    #     #     file_2 = None
    #     #     for filename in os.listdir(folder):
    #     #         # Проверка, что файл содержит search_term, но не содержит 'info'
    #     #         if 'info' not in filename and search_term in filename:
    #     #             if file_1 is None:
    #     #                 file_1 = os.path.join(folder, filename)
    #     #             elif file_2 is None:
    #     #                 file_2 = os.path.join(folder, filename)
    #     #     return file_1, file_2

    #     def find_file_in_folder(folder, search_term):
    #         file_1 = None
    #         file_2 = None
    #         files = os.listdir(folder)  # Список файлов в папке

    #         for filename in files:
    #             # Преобразуем имя файла в нижний регистр и проверяем условия
    #             if 'info' not in filename and search_term in filename:
    #                 if file_1 is None:
    #                     file_1 = os.path.join(folder, filename)
    #                 elif file_2 is None:
    #                     file_2 = os.path.join(folder, filename)
    #         return file_1, file_2

    #     # Найдите файл, содержащий search_string
    #     modified_attachment_path_1, modified_attachment_path_2 = find_file_in_folder(folder_path, search_string)

    #     if not modified_attachment_path_1 and not modified_attachment_path_2:
    #         # print(f"Файл, содержащий '{search_string}' в имени, не найден.")
    #         return

    #     # from datetime import datetime

    #     # # Получаем текущую дату и время
    #     # now = datetime.now()

    #     # # Форматируем дату и время в удобный вид
    #     # formatted_date_time = now.strftime("%Y-%m-%d %H:%M:%S")

    #     # import re
    #     # def extract_phone_number(text):
    #     #     # Регулярное выражение для поиска номера телефона
    #     #     phone_number_pattern = r'\b\d{11}\b'
    #     #     match = re.search(phone_number_pattern, text)

    #     #     if match:
    #     #         return match.group()
    #     #     return None

    #     # Настройки почтового сервера
    #     smtp_server = "smtp.mfc.tomsk.ru"
    #     smtp_port = 587  # Порт для TLS
    #     sender_email_address = "neverov@mfc.tomsk.ru"
    #     sender_password = "85HtnPdj"  # Убедитесь, что это пароль приложения

    #     from email.mime.multipart import MIMEMultipart

    #     # Создайте сообщение электронной почты
    #     # message = EmailMessage()
    #     message = MIMEMultipart()
    #     message["Subject"] = f"Документы Регистрация и Кадастрвоый учёт от {file_text.split(',')[0]}"
    #     message["From"] = sender_email_address
    #     message["To"] = 'msptosp@mfc.tomsk.ru'  # Адрес получателя

    #     # Установите текстовое содержимое письма
    #     # message.set_content(file_text)
    #     from email.mime.text import MIMEText
    #     message.attach(MIMEText(file_text, 'plain'))

    #     # Определение MIME-типов для произвольных файлов
    #     def guess_mime_type(filepath):
    #         import mimetypes
    #         mimetype, _ = mimetypes.guess_type(filepath)
    #         return mimetype or 'application/octet-stream'

    #     # Функция для прикрепления файла
    #     def attach_file(message, file_path):
    #         mime_type = guess_mime_type(file_path)
    #         maintype, subtype = mime_type.split('/', 1)
    #         with open(file_path, 'rb') as attachment:
    #             part = MIMEBase(maintype, subtype)
    #             part.set_payload(attachment.read())
    #             encoders.encode_base64(part)
    #             part.add_header('Content-Disposition', 'attachment', filename=os.path.basename(file_path))
    #             message.attach(part)

    #     attach_file(message, modified_attachment_path_1)
    #     attach_file(message, modified_attachment_path_2)

    #     # Функция для отправки письма
    #     def send_email():
    #         condition = False
    #         try:
    #             # Подключитесь к SMTP-серверу
    #             with smtplib.SMTP(smtp_server, smtp_port) as server:
    #                 server.starttls()  # Инициализация TLS
    #                 server.login(sender_email_address, sender_password)
    #                 server.send_message(message)  # Отправьте сообщение
    #             print("Письмо успешно отправлено.")

    #             # Удалите документ после отправки
    #             os.remove(modified_attachment_path_1)
    #             os.remove(modified_attachment_path_2)
    #             os.remove(file_with_info)

    #             condition = True
    #         except smtplib.SMTPAuthenticationError as e:
    #             print(f"Ошибка аутентификации mail: {e}")
    #         except Exception as e:
    #             print(f"Произошла ошибка mail: {e}")

    #         return condition

    #     # Попробуйте отправить письмо несколько раз с интервалами
    #     attempts = 3
    #     for attempt in range(attempts):
    #         print(f"Попытка {attempt + 1} из {attempts}")

    #         condition = send_email()
    #         if condition:
    #             break
    #         elif attempt < attempts - 1:
    #             print("Ожидание перед повторной попыткой...")
    #             time.sleep(10)  # Ожидание 10 секунд перед следующей попыткой

    #     return

    # except Exception as e:
    #     # Вывод подробной информации об ошибке
    #     print(f"Поймано исключение: {type(e).__name__}")
    #     print(f"Сообщение об ошибке: {str(e)}")
    #     import traceback
    #     print("Трассировка стека (stack trace):")
    #     traceback.print_exc()

    import os
    import requests

    # Ваш токен бота
    TOKEN = "6910991480:AAHecZctB4SuhT6AWfPbPb7SuxH97mONzN0"

    def search_user_id():
        import mysql.connector
        try:
            mydb = mysql.connector.connect(
                host="172.18.11.103",
                user="root",
                password="enigma1418",
                database="mdtomskbot",
                connect_timeout=2
            )
            mycursor = mydb.cursor()

            # Search in the database using ani and file_id
            query = f"SELECT id_tb, phone FROM file_id_tb WHERE marker = 'yes';"
            mycursor.execute(query)
            result = mycursor.fetchone()

            mydb.close()

            if not result:
                return None, None

            # Возвращаем результат
            return result[0], result[1]
        except Exception as e:
            print('search_in_base:', e)
            return {"code": "no"}

    def search_files(user_id):
        import mysql.connector
        try:
            mydb = mysql.connector.connect(
                host="172.18.11.103",
                user="root",
                password="enigma1418",
                database="mdtomskbot",
                connect_timeout=2
            )
            mycursor = mydb.cursor()

            # Search in the database using ani and file_id
            query = f"SELECT file, file_format, message FROM file_id_tb WHERE id_tb = '{user_id}' AND message IN ('salesman_passport', 'buyer_passport', 'proxy_passport', 'EGRN', 'right');"

            mycursor.execute(query)
            result = mycursor.fetchall()

            mydb.close()

            return result if result else None

        except Exception as e:
            print('search_in_base:', e)
            return {"code": "no"}

    def search_message(user_id):
        import mysql.connector
        try:
            mydb = mysql.connector.connect(
                host="172.18.11.103",
                user="root",
                password="enigma1418",
                database="mdtomskbot",
                connect_timeout=2
            )
            mycursor = mydb.cursor()

            # Search in the database using ani and file_id
            query = f"SELECT message FROM file_id_tb WHERE id_tb = '{user_id}' AND message LIKE '%Цена%';"
            mycursor.execute(query)
            result = mycursor.fetchall()

            mydb.close()

            return result[0][0] if result else None

        except Exception as e:
            print('search_in_base:', e)
            return {"code": "no"}

    def delete_files(user_id):
        import mysql.connector
        try:
            mydb = mysql.connector.connect(
                host="172.18.11.103",
                user="root",
                password="enigma1418",
                database="mdtomskbot",
                connect_timeout=2
            )
            mycursor = mydb.cursor()

            # Search in the database using ani and file_id
            query = f"DELETE FROM file_id_tb WHERE id_tb = '{user_id}';"
            mycursor.execute(query)
            mydb.commit()

            mydb.close()

            return

        except Exception as e:
            print('search_in_base:', e)
            return {"code": "no"}

    def send_files():
        # while True:
        try:
            user_id, phone = search_user_id()

            if user_id and phone:
                files = search_files(user_id)

                if files:
                    for file in files:

                        file_name = f'{file[2]}_{user_id}.{file[1]}'

                        # elif "photo" in message:
                        #     # Выбираем последнее (самое большое) фото
                        #     largest_photo = message["photo"][-1]
                        #     file_id = largest_photo["file_id"]
                        #     file_name = f"{file_id}.jpg"
                        #     print(f"file_id фотографии: {file_id}")

                        # elif "video" in message:
                        #     file_id = message["video"]["file_id"]
                        #     file_name = f"{file_id}.mp4"
                        #     print(f"file_id видео: {file_id}")

                        # Получаем путь к файлу
                        file_info_url = f"https://api.telegram.org/bot{TOKEN}/getFile?file_id={file[0]}"
                        file_info_response = requests.get(file_info_url)
                        file_info = file_info_response.json()

                        if file_info.get("ok"):
                            file_path = file_info['result']['file_path']

                            # Формируем URL для загрузки файла
                            file_url = f"https://api.telegram.org/file/bot{TOKEN}/{file_path}"

                            # Загружаем файл
                            file_response = requests.get(file_url)

                            # Создаем папку "321", если она не существует
                            save_directory = "C:\\Users\\admin\\Desktop\\file"
                            os.makedirs(save_directory, exist_ok=True)

                            # Сохраняем файл локально в папке "321"
                            save_path = os.path.join(save_directory, file_name)
                            with open(save_path, "wb") as f:
                                f.write(file_response.content)

                            print(f"Файл сохранен по пути: {save_path}")
                        else:
                            print("Ошибка при получении пути к файлу:", file_info)

                    message = search_message(user_id)
                    # Разбиваем строку по символу "_"
                    if message:
                        split_text = message.split('_')
                        # Соединяем элементы попарно с ": "
                        result = [f"{split_text[i]} {split_text[i+1]}" for i in range(0, len(split_text)-1, 2)]
                        # print(result)
                        result.append(phone)
                        result = ', '.join(result)

                        file_path = os.path.join(save_directory, f'info_{user_id}.txt')  # Путь к файлу

                        # Записываем file_path в файл
                        with open(file_path, 'w') as file:
                            file.write(result)

                    delete_files(user_id)

            import smtplib
            import time
            from email.mime.base import MIMEBase
            from email import encoders

            # Путь к папке, где находятся файлы
            folder_path = "C:\\Users\\admin\\Desktop\\file"

            # Функция для поиска первого файла, содержащего "info" в имени
            def find_file_with_info(folder):
                for filename in os.listdir(folder):
                    if f"info" in filename:
                        return os.path.join(folder, filename), str(filename).split('_')[1].split('.')[0]
                return None, None

            # Найдите файл с "info" в имени
            file_with_info, user_id = find_file_with_info(folder_path)

            if not file_with_info:
                # print("Файл, содержащий 'info', не найден.")
                return

            # Функция для извлечения текста из файла
            def extract_text_from_file(filepath):
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        return f.read()
                except:
                    with open(filepath, 'r', encoding='windows-1251') as f:
                        return f.read()

            # Извлекаем текст из файла
            file_text = extract_text_from_file(file_with_info)

            # Функция для получения первой части имени файла до символа "_"
            def extract_search_string_from_filename(found_files):
                search_strings = []
                for file in found_files:
                    search_strings.append(file.split("\\")[-1].split('.')[0] if "_" in file else None)
                return search_strings if search_strings else None

            # Функция для поиска первого файла, не содержащего "info" в имени
            def get_files_in_folder(folder,user_id):
                found_files = []
                files = os.listdir(folder)
                for filename in files:
                    if not "info" in filename and user_id in filename:
                        found_files.append(os.path.join(folder, filename))
                return found_files if found_files else None

            # Найдите первый файл в папке
            found_files = get_files_in_folder(folder_path, user_id)

            if not found_files:
                # print("В папке нет файлов.")
                return

            # Извлекаем строку для поиска из имени файла
            search_strings = extract_search_string_from_filename(found_files)

            if not search_strings:
                # print(f"Не удалось извлечь строку из имени файла: {found_files}")
                return

            def find_file_in_folder(folder, search_strings):
                files_found = []
                files = os.listdir(folder)  # Список файлов в папке

                for filename in files:
                    # Преобразуем имя файла в нижний регистр и проверяем условия
                    if 'info' not in filename and str(filename).split('.')[0] in search_strings:
                        files_found.append(os.path.join(folder, filename))

                return files_found if files_found else None

            # Найдите файл, содержащий search_string
            modified_attachment_path = find_file_in_folder(folder_path, search_strings)

            if not modified_attachment_path:
                # print(f"Файл, содержащий '{search_strings}' в имени, не найден.")
                return

            # Настройки почтового сервера
            smtp_server = "smtp.mfc.tomsk.ru"
            smtp_port = 587  # Порт для TLS
            sender_email_address = "neverov@mfc.tomsk.ru"
            sender_password = "85HtnPdj"  # Убедитесь, что это пароль приложения

            from email.mime.multipart import MIMEMultipart

            # Создайте сообщение электронной почты
            # message = EmailMessage()
            message = MIMEMultipart()
            message["Subject"] = f"Документы на составление договора от {file_text.split(',')[-1]}"
            message["From"] = sender_email_address
            # message["To"] = 'msptosp@mfc.tomsk.ru'  # Адрес получателя
            message["To"] = 'gulyaeva@mfc.tomsk.ru'  # Адрес получателя

            # Установите текстовое содержимое письма
            # message.set_content(file_text)
            from email.mime.text import MIMEText
            message.attach(MIMEText(file_text, 'plain'))

            # Определение MIME-типов для произвольных файлов
            def guess_mime_type(filepath):
                import mimetypes
                mimetype, _ = mimetypes.guess_type(filepath)
                return mimetype or 'application/octet-stream'

            # Функция для прикрепления файла
            def attach_file(message, file_path):
                for file in file_path:
                    mime_type = guess_mime_type(file)
                    maintype, subtype = mime_type.split('/', 1)
                    with open(file, 'rb') as attachment:
                        part = MIMEBase(maintype, subtype)
                        part.set_payload(attachment.read())
                        encoders.encode_base64(part)
                        part.add_header('Content-Disposition', 'attachment', filename=os.path.basename(file))
                        message.attach(part)

            attach_file(message, modified_attachment_path)

            # Функция для отправки письма
            def send_email():
                condition = False
                try:
                    # Подключитесь к SMTP-серверу
                    with smtplib.SMTP(smtp_server, smtp_port) as server:
                        server.starttls()  # Инициализация TLS
                        server.login(sender_email_address, sender_password)
                        server.send_message(message)  # Отправьте сообщение
                    print("Письмо успешно отправлено.")

                    # Удалите документ после отправки
                    for file in modified_attachment_path:
                        os.remove(file)
                    os.remove(file_with_info)

                    condition = True
                except smtplib.SMTPAuthenticationError as e:
                    print(f"Ошибка аутентификации mail: {e}")
                except Exception as e:
                    print(f"Произошла ошибка mail: {e}")

                return condition

            # Попробуйте отправить письмо несколько раз с интервалами
            attempts = 3
            for attempt in range(attempts):
                print(f"Попытка {attempt + 1} из {attempts}")

                condition = send_email()
                if condition:
                    break
                elif attempt < attempts - 1:
                    print("Ожидание перед повторной попыткой...")
                    time.sleep(10)  # Ожидание 10 секунд перед следующей попыткой

            return

        except Exception as e:
            # Вывод подробной информации об ошибке
            print(f"Поймано исключение: {type(e).__name__}")
            print(f"Сообщение об ошибке: {str(e)}")
            import traceback
            print("Трассировка стека (stack trace):")
            traceback.print_exc()

    send_files()

# if __name__ == "__main__":

#     # создание и запуск потоков для каждого процесса
#     t1 = threading.Thread(target=process_mail)

#     while True:
#         if not t1.is_alive():
#             t1 = threading.Thread(target=process_mail)
#             t1.start()
